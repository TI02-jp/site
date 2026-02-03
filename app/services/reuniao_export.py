"""Geração de PDF de reuniões com cabeçalho timbrado.

Usa o template DOCX fornecido em static/models/ata-template.docx para
preservar o cabeçalho/rodapé e converte o conteúdo das decisões (HTML) em PDF.
"""

from __future__ import annotations

import os
import re
import tempfile
import html as html_lib
from datetime import datetime
from typing import Iterable
from pathlib import Path
from zipfile import ZipFile

from bs4 import BeautifulSoup
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert as docx2pdf_convert
from flask import current_app
from fpdf import FPDF, HTMLMixin

from app.models.tables import ClienteReuniao, User


def _slugify(value: str) -> str:
    """Return a filesystem-friendly slug."""
    return re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-") or "reuniao"


def _get_temp_dir() -> Path:
    """Get a reliable temporary directory with proper permissions.
    
    For Windows services, ensures the temp directory is accessible
    and has proper write permissions.
    """
    # Try instance folder first (usually has proper permissions in production)
    instance_dir = Path(current_app.instance_path)
    if instance_dir.exists():
        temp_dir = instance_dir / "tmp"
        try:
            temp_dir.mkdir(parents=True, exist_ok=True)
            # Test write permission
            test_file = temp_dir / ".test_write"
            test_file.touch()
            test_file.unlink()
            return temp_dir
        except (PermissionError, OSError):
            pass
    
    # Fall back to system temp directory
    try:
        temp_root = Path(tempfile.gettempdir())
        # For Windows services, create a subdirectory in temp
        temp_dir = temp_root / "flask_reuniao_pdf"
        temp_dir.mkdir(parents=True, exist_ok=True)
        return temp_dir
    except (PermissionError, OSError):
        # Last resort: use current working directory
        return Path.cwd()


class _HTMLPDF(FPDF, HTMLMixin):
    """Minimal HTML-capable PDF used as a fallback when docx2pdf is unavailable."""


def _add_html_to_docx(doc: Document, html_content: str) -> None:
    """Converte HTML do Quill para elementos do python-docx preservando formatação."""
    soup = BeautifulSoup(html_content, "html.parser")

    def get_indent_level(element):
        """Extrai nível de indentação das classes Quill."""
        classes = element.get("class", [])
        if isinstance(classes, str):
            classes = classes.split()
        for cls in classes:
            if cls.startswith("ql-indent-"):
                try:
                    return int(cls.replace("ql-indent-", ""))
                except ValueError:
                    pass
        return 0

    def get_alignment(element):
        """Extrai alinhamento das classes Quill."""
        classes = element.get("class", [])
        if isinstance(classes, str):
            classes = classes.split()
        for cls in classes:
            if cls == "ql-align-center":
                return WD_ALIGN_PARAGRAPH.CENTER
            elif cls == "ql-align-right":
                return WD_ALIGN_PARAGRAPH.RIGHT
            elif cls == "ql-align-justify":
                return WD_ALIGN_PARAGRAPH.JUSTIFY
        return WD_ALIGN_PARAGRAPH.LEFT

    def add_formatted_text(paragraph, element):
        """Adiciona texto formatado (negrito, itálico, etc) ao parágrafo."""
        if isinstance(element, str):
            text = element
            if text.strip():
                paragraph.add_run(text)
            return

        # Processa o elemento HTML
        if element.name in ["strong", "b"]:
            for child in element.children:
                if isinstance(child, str):
                    run = paragraph.add_run(child)
                    run.bold = True
                else:
                    add_formatted_text(paragraph, child)
        elif element.name in ["em", "i"]:
            for child in element.children:
                if isinstance(child, str):
                    run = paragraph.add_run(child)
                    run.italic = True
                else:
                    add_formatted_text(paragraph, child)
        elif element.name == "u":
            for child in element.children:
                if isinstance(child, str):
                    run = paragraph.add_run(child)
                    run.underline = True
                else:
                    add_formatted_text(paragraph, child)
        elif element.name == "br":
            paragraph.add_run("\n")
        elif element.name == "span":
            # Spans podem ter estilos inline
            for child in element.children:
                add_formatted_text(paragraph, child)
        else:
            # Para outros elementos, extrai o texto
            text = element.get_text()
            if text.strip():
                paragraph.add_run(text)

    def process_list(list_element, is_ordered=False, parent_indent=0):
        """Processa listas <ul> ou <ol> com suporte a aninhamento."""
        list_items = list_element.find_all("li", recursive=False)
        for idx, li in enumerate(list_items):
            # Extrai classe de indentação do item
            indent_level = get_indent_level(li)
            # Calcula indentação total: indentação do item + indentação do pai
            total_indent = parent_indent + indent_level
            alignment = get_alignment(li)

            # Cria parágrafo com marcador
            p = doc.add_paragraph()
            p.alignment = alignment
            
            # Aplica indentação em polegadas
            # Usa 0.5" por nível para ficar mais visível
            indent_inches = 0.5 * (total_indent + 1)
            p.paragraph_format.left_indent = Inches(indent_inches)

            # Adiciona marcador/número
            if is_ordered:
                prefix = f"{idx + 1}. "
            else:
                prefix = "• "

            p.add_run(prefix)

            # Processa conteúdo do item (texto e formatação)
            has_nested_list = False
            for child in li.children:
                if hasattr(child, "name") and child.name in ["ul", "ol"]:
                    has_nested_list = True
                else:
                    add_formatted_text(p, child)
            
            # Se houver listas aninhadas, processa-as com indentação aumentada
            for child in li.children:
                if hasattr(child, "name") and child.name in ["ul", "ol"]:
                    process_list(child, child.name == "ol", parent_indent=total_indent + 1)

    def process_element(element):
        """Processa um elemento HTML genérico."""
        if isinstance(element, str):
            text = element.strip()
            if text:
                doc.add_paragraph(text)
            return

        if element.name == "p":
            p = doc.add_paragraph()
            p.alignment = get_alignment(element)
            indent = get_indent_level(element)
            if indent > 0:
                p.paragraph_format.left_indent = Inches(0.25 * indent)

            for child in element.children:
                add_formatted_text(p, child)

        elif element.name == "ul":
            process_list(element, is_ordered=False)
        elif element.name == "ol":
            process_list(element, is_ordered=True)
        elif element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(element.name[1])
            doc.add_heading(element.get_text(), level=level)
        elif element.name == "br":
            doc.add_paragraph()
        elif element.name in ["div", "section", "article"]:
            # Processa filhos de containers
            for child in element.children:
                process_element(child)
        else:
            # Para outros elementos, tenta extrair texto
            text = element.get_text().strip()
            if text:
                doc.add_paragraph(text)

    # Processa todos os elementos filhos do body ou do soup diretamente
    body = soup.find("body")
    root = body if body else soup

    for child in root.children:
        if hasattr(child, "name") or (isinstance(child, str) and child.strip()):
            process_element(child)



def _truncate_decisoes_html(html_content: str, max_chars: int) -> tuple[str, bool]:
    """Return sanitized HTML truncated by plain-text length when needed."""
    if max_chars <= 0:
        return html_content, False
    try:
        text_only = BeautifulSoup(html_content or "", "html.parser").get_text("\n")
    except Exception:
        text_only = str(html_content or "")
    if len(text_only) <= max_chars:
        return html_content, False
    trimmed = text_only[:max_chars].rstrip()
    safe_text = html_lib.escape(trimmed)
    notice = "<p><em>Conteudo truncado por limite de tamanho.</em></p>"
    return f"<p>{safe_text}...</p>{notice}", True

def _resolve_participantes_labels(participantes_raw: Iterable | None) -> list[str]:
    """Return all participant names (users + guests) without duplicates."""

    if not participantes_raw:
        return []

    user_ids: list[int] = []
    guest_names: list[str] = []

    def _add_guest(name: str) -> None:
        cleaned = (name or "").strip()
        if cleaned:
            guest_names.append(cleaned[:255])

    # First pass: separate user IDs and guest names
    for participante in participantes_raw:
        if isinstance(participante, dict):
            p_type = (participante.get("type") or participante.get("tipo") or participante.get("kind") or "").lower()
            if p_type == "guest":
                _add_guest(participante.get("name") or participante.get("nome") or participante.get("label") or "")
                continue
            pid = participante.get("id") or participante.get("user_id")
            if isinstance(pid, int):
                user_ids.append(pid)
                continue
            _add_guest(participante.get("name") or participante.get("label") or "")
        elif isinstance(participante, int):
            user_ids.append(participante)
        elif isinstance(participante, str):
            _add_guest(participante)

    # Build user lookup to resolve names
    user_lookup: dict[int, User] = {}
    if user_ids:
        usuarios = User.query.filter(User.id.in_(user_ids)).all()
        user_lookup = {usuario.id: usuario for usuario in usuarios}

    resolved: list[str] = []
    seen: set[str] = set()

    def _add(name: str) -> None:
        cleaned = (name or "").strip()
        if cleaned and cleaned not in seen:
            resolved.append(cleaned)
            seen.add(cleaned)

    # Add resolved users in the order they appeared
    for pid in user_ids:
        usuario = user_lookup.get(pid)
        nome_usuario = None
        if usuario is not None:
            nome_usuario = getattr(usuario, "name", None) or getattr(usuario, "username", None)
        _add(nome_usuario or f"Usuario #{pid}")

    # Add guests
    for nome in guest_names:
        _add(nome)

    return resolved


def export_reuniao_decisoes_pdf(reuniao: ClienteReuniao) -> tuple[bytes, str]:
    """Render decisions into the timbrado template and return PDF bytes + filename."""
    # Get template path with absolute resolution for service compatibility
    template_filename = "ata-template.docx"
    template_path = Path(current_app.root_path) / "static" / "models" / template_filename

    template_available = True

    # Verify template exists + is readable (do not hard-fail; fallback will be used)
    if not template_path.exists():
        template_available = False
        current_app.logger.warning(f"Template file not found: {template_path} (using fallback)")
    elif not os.access(template_path, os.R_OK):
        template_available = False
        current_app.logger.warning(f"No read permission for template: {template_path} (using fallback)")

    # Load DOCX directly (no need to materialize from DOTX)
    try:
        doc = Document(str(template_path)) if template_available else Document()
    except PackageNotFoundError as exc:
        template_available = False
        current_app.logger.warning(f"Template file could not be opened: {template_path} ({exc}); using fallback")
        doc = Document()

    empresa_nome = getattr(reuniao.empresa, "nome_empresa", "") or "Empresa"
    data_str = reuniao.data.strftime("%d/%m/%Y") if getattr(reuniao, "data", None) else "Não informado"
    setor_nome = getattr(getattr(reuniao, "setor", None), "nome", None) or "Não informado"

    # ==== TÍTULO PRINCIPAL ====
    titulo = doc.add_heading("REUNIÃO DE ALINHAMENTO", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(31, 78, 120)  # Azul corporativo

    # Adiciona espaço após título
    doc.add_paragraph()

    # ==== SEÇÃO: IDENTIFICAÇÃO ====
    identificacao_heading = doc.add_heading("IDENTIFICAÇÃO", level=2)
    for run in identificacao_heading.runs:
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Empresa
    p_empresa = doc.add_paragraph()
    p_empresa.add_run(empresa_nome).bold = True

    # Data
    p_data = doc.add_paragraph()
    p_data.add_run(f"Data: {data_str}")

    # Setor
    p_setor = doc.add_paragraph()
    p_setor.add_run(f"Setor: {setor_nome}")

    # Adiciona espaço entre seções
    doc.add_paragraph()

    # ==== SEÇÃO: PARTICIPANTES ====
    participantes_heading = doc.add_heading("PARTICIPANTES", level=2)
    for run in participantes_heading.runs:
        run.font.color.rgb = RGBColor(31, 78, 120)

    # Monta lista completa de participantes vindos do campo da reunião
    todos_participantes = []
    seen = set()

    # Apenas quem está no campo de participantes (usuários + convidados)
    for nome in _resolve_participantes_labels(reuniao.participantes):
        if nome not in seen:
            todos_participantes.append(nome)
            seen.add(nome)

    # Renderiza a lista
    if todos_participantes:
        for nome in todos_participantes:
            p = doc.add_paragraph()
            p.add_run(nome)
    else:
        doc.add_paragraph("Não informado.")

    # Adiciona espaço entre seções
    doc.add_paragraph()
    doc.add_paragraph()  # Espaço extra entre Participantes e Assuntos Tratados

    # ==== SEÇÃO: ASSUNTOS TRATADOS ====
    assuntos_heading = doc.add_heading("ASSUNTOS TRATADOS", level=2)
    for run in assuntos_heading.runs:
        run.font.color.rgb = RGBColor(31, 78, 120)

    decisoes_html = reuniao.decisoes or "<p>Sem assuntos registrados.</p>"

    max_chars = int(os.getenv("PDF_DECISOES_MAX_CHARS", "100000"))
    decisoes_html, was_truncated = _truncate_decisoes_html(decisoes_html, max_chars)
    if was_truncated:
        current_app.logger.warning(f"[PDF Export] Decisoes truncadas para {max_chars} caracteres")

    # Usa a função que preserva a formatação do Quill
    try:
        _add_html_to_docx(doc, decisoes_html)
    except Exception as exc:
        current_app.logger.warning(f"Falha ao renderizar HTML no DOCX (fallback texto simples): {exc}")
        plain_text = BeautifulSoup(decisoes_html, "html.parser").get_text("\n").strip()
        if plain_text:
            doc.add_paragraph(plain_text)

    # Initialize variables for cleanup in finally block
    pdf_bytes: bytes | None = None
    tmp_docx_path: Path | None = None
    tmp_pdf_path: Path | None = None
    co_initialized = False
    pythoncom_mod = None
    temp_dir: Path | None = None

    # Wrap all temp file operations in try-except to trigger fallback on permission errors
    try:
        if os.getenv("PDF_SKIP_DOCX2PDF", "0") == "1":
            current_app.logger.info("[PDF Export] docx2pdf skipped by PDF_SKIP_DOCX2PDF=1 (using fallback)")
            raise RuntimeError("docx2pdf skipped by PDF_SKIP_DOCX2PDF=1")
        temp_dir = _get_temp_dir()
        current_app.logger.info(f"[PDF Export] Using temp directory: {temp_dir}")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False, dir=str(temp_dir)) as tmp_docx:
            doc.save(tmp_docx.name)
            tmp_docx_path = Path(tmp_docx.name)

        current_app.logger.info(f"[PDF Export] Created temp DOCX at: {tmp_docx_path}")

        try:
            import pythoncom as pythoncom_mod  # type: ignore

            pythoncom_mod.CoInitialize()
            co_initialized = True
            current_app.logger.info("[PDF Export] pythoncom.CoInitialize() succeeded")
        except Exception as pythoncom_error:
            # Se ja estiver inicializado ou indisponivel (ex.: Linux), seguimos
            pythoncom_mod = None
            current_app.logger.warning(f"[PDF Export] pythoncom initialization failed or not available: {pythoncom_error}")

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False, dir=str(temp_dir)) as tmp_pdf:
            tmp_pdf_path = Path(tmp_pdf.name)

        current_app.logger.info(f"[PDF Export] Starting docx2pdf conversion: {tmp_docx_path} -> {tmp_pdf_path}")

        # docx2pdf gera PDF usando Word/LibreOffice disponiveis
        try:
            docx2pdf_convert(str(tmp_docx_path), str(tmp_pdf_path))
            pdf_bytes = tmp_pdf_path.read_bytes()
            pdf_size_kb = len(pdf_bytes) / 1024
            current_app.logger.info(f"[PDF Export] ✓ PDF generated successfully using docx2pdf ({pdf_size_kb:.1f} KB)")
        except Exception as docx2pdf_error:
            current_app.logger.error(
                f"[PDF Export] ✗ docx2pdf conversion FAILED (will use fallback): {type(docx2pdf_error).__name__}: {docx2pdf_error}",
                exc_info=docx2pdf_error
            )
            raise
    except Exception as exc:
        # Log detailed error information for debugging production issues
        error_context = {
            "temp_dir": str(temp_dir) if temp_dir else "not_created",
            "tmp_docx_path": str(tmp_docx_path) if tmp_docx_path else "not_created",
            "error_type": type(exc).__name__
        }
        current_app.logger.warning(
            "[PDF Export] ⚠ Temp file operations or docx2pdf failed (context: %s); using fallback HTML->PDF renderer: %s",
            error_context, exc, exc_info=exc
        )

        # Use fallback PDF generation with robust error handling
        participantes_labels = todos_participantes or []
        try:
            current_app.logger.info("[PDF Export] Attempting fallback PDF generation using fpdf2 (NOTE: template header will NOT be included)")
            pdf_bytes = _render_pdf_fallback(
                empresa_nome=empresa_nome,
                data_str=data_str,
                setor_nome=setor_nome,
                participantes=participantes_labels,
                decisoes_html=decisoes_html,
                template_available=template_available,
            )
            pdf_size_kb = len(pdf_bytes) / 1024
            current_app.logger.warning(f"[PDF Export] ✓ PDF generated using FALLBACK renderer ({pdf_size_kb:.1f} KB) - Template header/logo NOT included")
        except Exception as fallback_error:
            current_app.logger.error(
                "Fallback PDF renderer also failed: %s", fallback_error, exc_info=fallback_error
            )
            # Re-raise to trigger 500 error with proper logging
            raise RuntimeError(
                "Falha ao gerar PDF da reunião usando ambos os métodos (docx2pdf e fallback)"
            ) from fallback_error
    finally:
        if co_initialized and pythoncom_mod:
            try:
                pythoncom_mod.CoUninitialize()
            except Exception:
                pass
        if tmp_docx_path:
            try:
                tmp_docx_path.unlink(missing_ok=True)
            except Exception as e:
                current_app.logger.debug(f"Não foi possível remover arquivo temporário DOCX: {e}")
        if tmp_pdf_path:
            try:
                tmp_pdf_path.unlink(missing_ok=True)
            except Exception as e:
                current_app.logger.debug(f"Não foi possível remover arquivo temporário PDF: {e}")
    
    if not pdf_bytes:
        raise RuntimeError("Falha ao gerar PDF da reunião.")

    date_for_name = (
        reuniao.data.strftime("%Y-%m-%d") if getattr(reuniao, "data", None) else datetime.now().date()
    )
    empresa_slug = _slugify(empresa_nome)
    filename = f"reuniao-{empresa_slug}-{date_for_name}.pdf"

    return pdf_bytes, filename


def _render_pdf_fallback(
    *,
    empresa_nome: str,
    data_str: str,
    setor_nome: str,
    participantes: list[str],
    decisoes_html: str,
    template_available: bool = True,
) -> bytes:
    """Generate a lightweight PDF using fpdf2 when docx2pdf/Word is unavailable."""

    pdf = _HTMLPDF()
    unicode_font_loaded = False

    def _try_load_unicode_font() -> None:
        nonlocal unicode_font_loaded
        # Prefer a known Unicode font if available in the venv
        dejavu_path = Path(current_app.root_path) / "venv" / "Lib" / "site-packages" / "matplotlib" / "mpl-data" / "fonts" / "ttf" / "DejaVuSans.ttf"
        if dejavu_path.exists():
            try:
                pdf.add_font("DejaVuSans", "", str(dejavu_path), uni=True)
                pdf.set_font("DejaVuSans", "", 12)
                unicode_font_loaded = True
                current_app.logger.info(f"[PDF Export] Unicode font loaded for fallback: {dejavu_path}")
                return
            except Exception as exc:
                current_app.logger.warning(f"[PDF Export] Failed to load Unicode font {dejavu_path}: {exc}")

        # Keep default Helvetica if no Unicode font is available
        try:
            pdf.set_font("Helvetica", "", 12)
        except Exception:
            pass

    def _sanitize_text(text: str) -> str:
        # Remove BOM and unsupported characters when not using a Unicode font
        if not text:
            return text
        cleaned = text.replace("\ufeff", "")
        if unicode_font_loaded:
            return cleaned
        try:
            return cleaned.encode("latin-1").decode("latin-1")
        except UnicodeEncodeError:
            return cleaned.encode("latin-1", errors="replace").decode("latin-1")

    def _safe_write_html(html: str, fallback_text: str | None = None) -> None:
        try:
            safe_html = _sanitize_text(html)
            pdf.write_html(safe_html)
        except Exception as exc:
            current_app.logger.warning(f"Fallback HTML rendering failed (plain text used): {exc}")
            text_content = BeautifulSoup(html or "", "html.parser").get_text("\n").strip()
            final_text = text_content or (fallback_text or "")
            final_text = _sanitize_text(final_text)
            if final_text:
                if pdf.page == 0:
                    pdf.add_page()
                pdf.multi_cell(0, 6, final_text)

    # Extract header and footer images from template
    header_img_path = None
    footer_img_path = None

    if template_available:
        try:
            template_path = Path(current_app.root_path) / "static" / "models" / "ata-template.docx"
            if template_path.exists():
                from zipfile import ZipFile
                import tempfile

                # Extract images from DOCX (DOCX is a ZIP file)
                with ZipFile(template_path, 'r') as docx_zip:
                    # Look for images in word/media/
                    image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]

                    if len(image_files) >= 1:
                        # Extract header image (first image)
                        header_image = image_files[0]
                        header_data = docx_zip.read(header_image)

                        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(header_image).suffix) as tmp_img:
                            tmp_img.write(header_data)
                            header_img_path = tmp_img.name

                    if len(image_files) >= 2:
                        # Extract footer image (second image)
                        footer_image = image_files[1]
                        footer_data = docx_zip.read(footer_image)

                        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(footer_image).suffix) as tmp_img:
                            tmp_img.write(footer_data)
                            footer_img_path = tmp_img.name
        except Exception as img_error:
            current_app.logger.warning(f"Could not extract images from template: {img_error}")

    _try_load_unicode_font()

    # Set auto page break with margin for footer
    pdf.set_auto_page_break(auto=True, margin=40)  # Space for footer
    pdf.add_page()

    # Add header image at top
    if header_img_path:
        try:
            pdf.image(header_img_path, x=10, y=8, w=190)
            pdf.ln(40)  # Space after header image
        except Exception as e:
            current_app.logger.warning(f"Could not add header image: {e}")

    # Add footer image on every page
    if footer_img_path:
        def footer_with_image():
            pdf.set_y(-35)  # Position 35mm from bottom
            try:
                pdf.image(footer_img_path, x=10, y=pdf.get_y(), w=190)
            except Exception as e:
                current_app.logger.warning(f"Could not add footer image: {e}")

        # Override footer method
        pdf.footer = footer_with_image

    if unicode_font_loaded:
        pdf.set_font("DejaVuSans", "B", 16)
    else:
        pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "REUNIÃO DE ALINHAMENTO", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(4)
    if unicode_font_loaded:
        pdf.set_font("DejaVuSans", "", 12)
    else:
        pdf.set_font("Helvetica", "", 12)

    meta_html = (
        f"<p><b>Empresa:</b> {empresa_nome}<br>"
        f"<b>Data:</b> {data_str}<br>"
        f"<b>Setor:</b> {setor_nome}</p>"
    )
    _safe_write_html(meta_html)

    participantes_html = "<h3>Participantes</h3>"
    if participantes:
        participantes_html += "<ul>" + "".join(f"<li>{p}</li>" for p in participantes) + "</ul>"
    else:
        participantes_html += "<p>Não informado.</p>"
    _safe_write_html(participantes_html, fallback_text="Participantes")

    # Adiciona espaço extra entre Participantes e Assuntos Tratados
    pdf.ln(8)

    decisoes_section = "<h3></h3>" + (decisoes_html or "<p>Sem assuntos registrados.</p>")
    _safe_write_html(decisoes_section, fallback_text="Assuntos e decisoes")

    # pdf.output() retorna bytes diretamente no fpdf2 moderno
    pdf_output = pdf.output()

    # Clean up temporary image files
    if header_img_path:
        try:
            Path(header_img_path).unlink(missing_ok=True)
        except Exception:
            pass

    if footer_img_path:
        try:
            Path(footer_img_path).unlink(missing_ok=True)
        except Exception:
            pass

    # Se for string (versão antiga do fpdf), converte para bytes
    if isinstance(pdf_output, str):
        try:
            return pdf_output.encode("latin-1")
        except UnicodeEncodeError:
            current_app.logger.warning("Encoding PDF with replacements for incompatible characters")
            return pdf_output.encode("latin-1", errors="replace")

    # Já é bytes, retorna diretamente
    return pdf_output
